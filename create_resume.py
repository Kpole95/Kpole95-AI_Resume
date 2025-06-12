from docx import Document
import os

def create_resume():
    doc = Document()
    doc.add_heading("Murali Krishna Pole", 0)
    doc.add_paragraph("Hyderabad")

    doc.add_heading("Skills", 1)
    skills = ["Python", "SQL", "Docker", "Machine Learning", "HTML", "CSS", "Git", "Linux"]
    for skill in skills:
        doc.add_paragraph(f"- {skill}")

    doc.add_heading("Experience", 1)
    doc.add_paragraph("None")

    doc.add_heading("Education", 1)
    doc.add_paragraph("Bachelor of Technology in Computer Science")
    doc.add_paragraph("SKILLS: Technical")

    output_dir = "data/resumes_raw"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    output_path = os.path.join(output_dir, "Krishna_Pole_Resume.docx")
    doc.save(output_path)
    print(f"Resume saved to {output_path}")

if __name__ == "__main__":
    create_resume()
